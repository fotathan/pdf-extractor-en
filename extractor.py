import streamlit as st
import pdfplumber
import pandas as pd
import json
import requests
import zipfile
import time
import re
import random
from io import BytesIO
from pathlib import Path
from docx import Document


# ------------------------------------------------------------
# Custom exceptions
# ------------------------------------------------------------
class RateLimitError(Exception):
    pass


# ------------------------------------------------------------
# Page configuration
# ------------------------------------------------------------
st.set_page_config(page_title="AI Procurement Extractor", layout="wide")
st.title("🤖 AI-Powered Tender Data Extractor")

st.markdown(
    """
This tool extracts tender/product tables from uploaded files and converts them into structured outputs.

**Supported input**
- PDF
- XLSX / XLS
- DOCX
- CSV
- TXT / Markdown / HTML / JSON

**Supported output**
- Excel
- CSV
- HTML
- Markdown
- JSON

The app tries to use **direct parsing first** for already-tabular files like XLSX / XLS / CSV / JSON.
If needed, it falls back to AI extraction.

**Language behavior**
- Headers stay in the original source language
- Values stay in the original source language
- The app does not intentionally translate exported output into English

**Numeric handling**
- Original values are preserved
- Optional helper numeric columns can be added for easier sorting/filtering/analysis
"""
)

# ------------------------------------------------------------
# Sidebar
# ------------------------------------------------------------
st.sidebar.header("Configuration")

api_key = st.sidebar.text_input(
    "1. Gemini API Key",
    type="password",
    help="Enter your Google AI Studio Gemini API Key",
)

uploaded_files = st.file_uploader(
    "2. Upload one or more files",
    type=["pdf", "xlsx", "xls", "docx", "csv", "txt", "md", "markdown", "html", "htm", "json"],
    accept_multiple_files=True,
)

output_formats = st.sidebar.multiselect(
    "3. Output formats",
    options=["xlsx", "csv", "html", "md", "json"],
    default=["xlsx"],
)

chunk_size = st.sidebar.slider(
    "4. Chunk size (characters per AI request)",
    min_value=4000,
    max_value=30000,
    value=18000,
    step=1000,
    help="Larger chunks reduce request count. Too small a chunk size can increase 429 errors.",
)

timeout_seconds = st.sidebar.slider(
    "5. API timeout (seconds)",
    min_value=60,
    max_value=600,
    value=300,
    step=30,
)

max_retries = st.sidebar.slider(
    "6. Retries per chunk",
    min_value=0,
    max_value=5,
    value=2,
    step=1,
)

direct_parse_mode = st.sidebar.selectbox(
    "7. Structured file handling",
    options=[
        "Use direct parsing first",
        "Always use AI",
    ],
    index=0,
)

min_api_interval_seconds = st.sidebar.slider(
    "8. Minimum interval between Gemini calls (seconds)",
    min_value=0.0,
    max_value=30.0,
    value=12.0,
    step=1.0,
    help="Global pacing between Gemini requests. Helps reduce 429 errors much more effectively than short chunk delays.",
)

inter_file_delay = st.sidebar.slider(
    "9. Delay between files (seconds)",
    min_value=0.0,
    max_value=20.0,
    value=3.0,
    step=0.5,
)

add_numeric_helper_columns_flag = st.sidebar.checkbox(
    "10. Add normalized numeric helper columns",
    value=True,
    help="Keeps original values unchanged and adds extra helper columns for numeric-looking fields.",
)

numeric_detection_threshold = st.sidebar.slider(
    "11. Numeric detection threshold",
    min_value=0.3,
    max_value=1.0,
    value=0.6,
    step=0.05,
    help="A column gets a numeric helper version if this share of its non-empty values can be parsed as numbers.",
)

show_preview_before_processing = st.sidebar.checkbox(
    "12. Show sheet/table preview before processing",
    value=False,
    help="Preview Excel sheets and detected tables before processing files.",
)

preview_rows = st.sidebar.slider(
    "13. Preview rows",
    min_value=3,
    max_value=20,
    value=8,
    step=1,
)

run_button = st.sidebar.button("🚀 Process Files", type="primary")

# ------------------------------------------------------------
# Session state
# ------------------------------------------------------------
if "results" not in st.session_state:
    st.session_state.results = []

if "last_gemini_call_ts" not in st.session_state:
    st.session_state.last_gemini_call_ts = 0.0

if "preview_selections" not in st.session_state:
    st.session_state.preview_selections = {}


# ------------------------------------------------------------
# Utility helpers
# ------------------------------------------------------------
def safe_filename(name: str) -> str:
    stem = Path(name).stem
    stem = re.sub(r"[^\w\-.]+", "_", stem, flags=re.UNICODE)
    return stem[:120] if stem else "output"


def normalize_cell_value(value):
    if pd.isna(value):
        return ""
    return str(value).strip()


def preserve_original_columns(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()

    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]

    for col in df.columns:
        df[col] = df[col].map(normalize_cell_value)

    df = df.loc[~(df.apply(lambda row: all(v == "" for v in row), axis=1))].copy()
    df = df.drop_duplicates().reset_index(drop=True)
    return df


def merge_dfs(dfs):
    valid = [preserve_original_columns(df) for df in dfs if df is not None and not df.empty]
    if not valid:
        return pd.DataFrame()

    all_columns = []
    for df in valid:
        for c in df.columns:
            if c not in all_columns:
                all_columns.append(c)

    aligned = [df.reindex(columns=all_columns, fill_value="") for df in valid]
    merged = pd.concat(aligned, ignore_index=True)
    return preserve_original_columns(merged)


def clean_llm_json(text_response: str) -> str:
    text_response = text_response.strip()

    if text_response.startswith("```json"):
        text_response = text_response[7:]
    elif text_response.startswith("```"):
        text_response = text_response[3:]

    if text_response.endswith("```"):
        text_response = text_response[:-3]

    return text_response.strip()


def chunk_text(text: str, max_chars: int = 18000):
    if not text:
        return []

    lines = text.splitlines()
    chunks = []
    current = []
    current_len = 0

    for line in lines:
        line_len = len(line) + 1
        if current and current_len + line_len > max_chars:
            chunks.append("\n".join(current).strip())
            current = [line]
            current_len = line_len
        else:
            current.append(line)
            current_len += line_len

    if current:
        chunks.append("\n".join(current).strip())

    return [c for c in chunks if c.strip()]


def looks_like_useful_table(df: pd.DataFrame) -> bool:
    if df is None or df.empty:
        return False

    df = preserve_original_columns(df)

    if df.empty:
        return False

    rows, cols = df.shape
    if cols < 2:
        return False

    non_empty_cells = (df != "").sum().sum()
    if non_empty_cells < max(6, rows * 2):
        return False

    return rows >= 2 and cols >= 2


def enforce_min_api_interval(min_interval_seconds: float):
    now = time.time()
    last_call = st.session_state.get("last_gemini_call_ts", 0.0)
    elapsed = now - last_call

    if elapsed < min_interval_seconds:
        time.sleep(min_interval_seconds - elapsed)

    st.session_state["last_gemini_call_ts"] = time.time()


# ------------------------------------------------------------
# Numeric normalization helpers
# ------------------------------------------------------------
def parse_localized_number(value):
    if value is None:
        return None

    s = str(value).strip()
    if s == "":
        return None

    s = s.replace("\u00A0", " ").replace("\u202F", " ").strip()

    negative = False
    if s.startswith("(") and s.endswith(")"):
        negative = True
        s = s[1:-1].strip()

    s = re.sub(r"[^0-9,\.\-\' %]", "", s).strip()

    if s == "" or s in {"-", ".", ",", "%"}:
        return None

    percent = "%" in s
    s = s.replace("%", "").strip()
    s = s.replace(" ", "").replace("'", "")

    if s.count("-") > 1:
        return None
    if "-" in s and not s.startswith("-"):
        return None

    comma_count = s.count(",")
    dot_count = s.count(".")

    if comma_count > 0 and dot_count > 0:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "")
            s = s.replace(",", ".")
        else:
            s = s.replace(",", "")
    elif comma_count > 0:
        if comma_count > 1:
            parts = s.split(",")
            if all(len(p) == 3 for p in parts[1:] if p != ""):
                s = "".join(parts)
            else:
                s = "".join(parts[:-1]) + "." + parts[-1]
        else:
            left, right = s.split(",", 1)
            if right.isdigit() and len(right) == 3 and len(left.replace("-", "")) >= 1:
                s = left + right
            else:
                s = left + "." + right
    elif dot_count > 0:
        if dot_count > 1:
            parts = s.split(".")
            if all(len(p) == 3 for p in parts[1:] if p != ""):
                s = "".join(parts)
            else:
                s = "".join(parts[:-1]) + "." + parts[-1]
        else:
            left, right = s.split(".", 1)
            if right.isdigit() and len(right) == 3 and len(left.replace("-", "")) >= 1:
                s = left + right

    try:
        number = float(s)
    except Exception:
        return None

    if negative:
        number = -number

    if percent:
        return number / 100.0

    return number


def should_add_numeric_helper(series: pd.Series, threshold: float = 0.6) -> bool:
    values = [str(v).strip() for v in series if str(v).strip() != ""]
    if len(values) < 2:
        return False

    parsed = [parse_localized_number(v) for v in values]
    parsed_ok = sum(v is not None for v in parsed)
    ratio = parsed_ok / len(values) if values else 0.0

    return ratio >= threshold and parsed_ok >= 2


def add_numeric_helper_columns(df: pd.DataFrame, threshold: float = 0.6) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()

    df = preserve_original_columns(df)

    helper_columns = {}
    original_columns = list(df.columns)

    for col in original_columns:
        if str(col).endswith(" [num]"):
            continue

        series = df[col]
        if should_add_numeric_helper(series, threshold=threshold):
            helper_col_name = f"{col} [num]"
            parsed_values = [parse_localized_number(v) for v in series]
            helper_columns[helper_col_name] = pd.to_numeric(parsed_values, errors="coerce")

    if not helper_columns:
        return df

    for helper_name, helper_series in helper_columns.items():
        df[helper_name] = helper_series

    return df


def finalize_output_df(df: pd.DataFrame, add_numeric_helpers: bool, threshold: float) -> pd.DataFrame:
    df = preserve_original_columns(df)
    if add_numeric_helpers:
        df = add_numeric_helper_columns(df, threshold=threshold)
    return df


# ------------------------------------------------------------
# Preview helpers
# ------------------------------------------------------------
def get_excel_sheet_previews(uploaded_file, max_rows: int = 8):
    uploaded_file.seek(0)
    previews = []

    try:
        xls = pd.ExcelFile(uploaded_file)
    except Exception:
        return previews

    for sheet_name in xls.sheet_names:
        try:
            df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str).fillna("")
            df = preserve_original_columns(df)
            previews.append(
                {
                    "sheet_name": sheet_name,
                    "rows": len(df),
                    "cols": len(df.columns),
                    "looks_useful": looks_like_useful_table(df),
                    "preview_df": df.head(max_rows),
                }
            )
        except Exception as exc:
            previews.append(
                {
                    "sheet_name": sheet_name,
                    "rows": 0,
                    "cols": 0,
                    "looks_useful": False,
                    "preview_df": pd.DataFrame({"Error": [str(exc)]}),
                }
            )

    return previews


def get_pdf_table_previews(uploaded_file, max_rows: int = 8):
    uploaded_file.seek(0)
    previews = []

    try:
        with pdfplumber.open(uploaded_file) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    tables = page.extract_tables(
                        table_settings={
                            "vertical_strategy": "lines",
                            "horizontal_strategy": "lines",
                            "intersection_tolerance": 5,
                            "snap_tolerance": 3,
                            "join_tolerance": 3,
                            "edge_min_length": 3,
                            "min_words_vertical": 1,
                            "min_words_horizontal": 1,
                        }
                    )
                except Exception:
                    tables = []

                for table_index, table in enumerate(tables, start=1):
                    if not table:
                        continue

                    cleaned_rows = []
                    max_cols = 0

                    for row in table:
                        if not row:
                            continue
                        cleaned_row = []
                        for cell in row:
                            cleaned_row.append("" if cell is None else str(cell).replace("\n", " ").strip())
                        if any(cell != "" for cell in cleaned_row):
                            cleaned_rows.append(cleaned_row)
                            max_cols = max(max_cols, len(cleaned_row))

                    if len(cleaned_rows) < 2 or max_cols < 2:
                        continue

                    normalized_rows = [row + [""] * (max_cols - len(row)) for row in cleaned_rows]
                    preview_df = pd.DataFrame(normalized_rows[:max_rows])

                    previews.append(
                        {
                            "page_num": page_num,
                            "table_index": table_index,
                            "rows": len(cleaned_rows),
                            "cols": max_cols,
                            "preview_df": preview_df,
                        }
                    )
    except Exception:
        return previews

    return previews


def get_docx_table_previews(uploaded_file, max_rows: int = 8):
    uploaded_file.seek(0)
    previews = []

    try:
        doc = Document(uploaded_file)
    except Exception:
        return previews

    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        max_cols = 0
        for row in table.rows:
            row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cell != "" for cell in row_cells):
                rows.append(row_cells)
                max_cols = max(max_cols, len(row_cells))

        if len(rows) < 2 or max_cols < 2:
            continue

        normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]
        preview_df = pd.DataFrame(normalized_rows[:max_rows])

        previews.append(
            {
                "table_index": table_index,
                "rows": len(rows),
                "cols": max_cols,
                "preview_df": preview_df,
            }
        )

    return previews


def render_previews(uploaded_files, preview_rows: int):
    if not uploaded_files:
        return

    st.markdown("## Preview before processing")

    for uploaded_file in uploaded_files:
        suffix = Path(uploaded_file.name).suffix.lower()
        file_key = safe_filename(uploaded_file.name)

        with st.expander(f"Preview: {uploaded_file.name}", expanded=False):
            if suffix in [".xlsx", ".xls"]:
                sheet_previews = get_excel_sheet_previews(uploaded_file, max_rows=preview_rows)

                if not sheet_previews:
                    st.info("No sheet preview available.")
                    continue

                available_sheets = [p["sheet_name"] for p in sheet_previews]
                default_sheets = [p["sheet_name"] for p in sheet_previews if p["looks_useful"]]
                if not default_sheets:
                    default_sheets = available_sheets

                selection_key = f"sheet_select_{file_key}"
                selected_sheets = st.multiselect(
                    f"Sheets to process for {uploaded_file.name}",
                    options=available_sheets,
                    default=st.session_state.preview_selections.get(selection_key, default_sheets),
                    key=selection_key,
                )
                st.session_state.preview_selections[selection_key] = selected_sheets

                for p in sheet_previews:
                    badge = "✅ likely useful" if p["looks_useful"] else "⚪ preview only"
                    st.markdown(
                        f"**Sheet:** {p['sheet_name']} | Rows: {p['rows']} | Columns: {p['cols']} | {badge}"
                    )
                    st.dataframe(p["preview_df"], use_container_width=True)

            elif suffix == ".pdf":
                table_previews = get_pdf_table_previews(uploaded_file, max_rows=preview_rows)
                if not table_previews:
                    st.info("No detectable PDF tables found in preview. The file can still be processed via text extraction.")
                else:
                    st.markdown(f"Detected **{len(table_previews)}** PDF table(s).")
                    for p in table_previews:
                        st.markdown(
                            f"**Page {p['page_num']} - Table {p['table_index']}** | Rows: {p['rows']} | Columns: {p['cols']}"
                        )
                        st.dataframe(p["preview_df"], use_container_width=True)

            elif suffix == ".docx":
                table_previews = get_docx_table_previews(uploaded_file, max_rows=preview_rows)
                if not table_previews:
                    st.info("No DOCX tables found in preview.")
                else:
                    st.markdown(f"Detected **{len(table_previews)}** DOCX table(s).")
                    for p in table_previews:
                        st.markdown(
                            f"**Table {p['table_index']}** | Rows: {p['rows']} | Columns: {p['cols']}"
                        )
                        st.dataframe(p["preview_df"], use_container_width=True)

            elif suffix in [".csv", ".json"]:
                try:
                    if suffix == ".csv":
                        df = direct_parse_csv(uploaded_file)
                    else:
                        df = direct_parse_json(uploaded_file)

                    if df is None or df.empty:
                        st.info("No structured preview available.")
                    else:
                        st.markdown(f"Rows: **{len(df)}** | Columns: **{len(df.columns)}**")
                        st.dataframe(df.head(preview_rows), use_container_width=True)
                except Exception as exc:
                    st.error(f"Preview failed: {exc}")

            else:
                st.info("Preview not available for this file type.")


# ------------------------------------------------------------
# Export helpers
# ------------------------------------------------------------
def df_to_excel_bytes(df: pd.DataFrame) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Extracted Data")
    return output.getvalue()


def df_to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8-sig")


def df_to_html_bytes(df: pd.DataFrame) -> bytes:
    return df.to_html(index=False, border=1).encode("utf-8")


def df_to_md_bytes(df: pd.DataFrame) -> bytes:
    try:
        md = df.to_markdown(index=False)
    except Exception:
        md = df.to_csv(index=False)
    return md.encode("utf-8")


def df_to_json_bytes(df: pd.DataFrame) -> bytes:
    return df.to_json(orient="records", force_ascii=False, indent=2).encode("utf-8")


def get_export_bytes(df: pd.DataFrame, fmt: str) -> bytes:
    if fmt == "xlsx":
        return df_to_excel_bytes(df)
    if fmt == "csv":
        return df_to_csv_bytes(df)
    if fmt == "html":
        return df_to_html_bytes(df)
    if fmt == "md":
        return df_to_md_bytes(df)
    if fmt == "json":
        return df_to_json_bytes(df)
    raise ValueError(f"Unsupported format: {fmt}")


def build_zip(results, formats):
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zip_file:
        for result in results:
            if result["status"] != "success":
                continue

            df = result["df"]
            base_name = safe_filename(result["file_name"])

            for fmt in formats:
                try:
                    content = get_export_bytes(df, fmt)
                    zip_file.writestr(f"{base_name}.{fmt}", content)
                except Exception as exc:
                    zip_file.writestr(
                        f"{base_name}_{fmt}_ERROR.txt",
                        str(exc).encode("utf-8")
                    )

    zip_buffer.seek(0)
    return zip_buffer.getvalue()


# ------------------------------------------------------------
# Direct parsers (language-preserving)
# ------------------------------------------------------------
def direct_parse_csv(uploaded_file):
    uploaded_file.seek(0)

    encodings = ["utf-8", "utf-8-sig", "cp1252", "latin-1"]
    separators = [",", ";", "\t", "|"]

    raw = uploaded_file.read()

    for enc in encodings:
        for sep in separators:
            try:
                text = raw.decode(enc)
                df = pd.read_csv(BytesIO(text.encode("utf-8")), sep=sep, dtype=str).fillna("")
                df = preserve_original_columns(df)
                if looks_like_useful_table(df):
                    return df
            except Exception:
                continue

    return pd.DataFrame()


def direct_parse_json(uploaded_file):
    uploaded_file.seek(0)
    raw = uploaded_file.read().decode("utf-8", errors="ignore")

    try:
        parsed = json.loads(raw)
    except Exception:
        return pd.DataFrame()

    if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
        df = pd.DataFrame(parsed)
        df = preserve_original_columns(df)
        return df if looks_like_useful_table(df) else pd.DataFrame()

    if isinstance(parsed, dict):
        for value in parsed.values():
            if isinstance(value, list) and value and isinstance(value[0], dict):
                df = pd.DataFrame(value)
                df = preserve_original_columns(df)
                if looks_like_useful_table(df):
                    return df

    return pd.DataFrame()


def direct_parse_excel(uploaded_file, selected_sheets=None):
    uploaded_file.seek(0)
    xls = pd.ExcelFile(uploaded_file)
    dfs = []

    sheet_names = selected_sheets if selected_sheets else xls.sheet_names

    for sheet_name in sheet_names:
        try:
            df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str).fillna("")
            df = preserve_original_columns(df)
            if looks_like_useful_table(df):
                if len(sheet_names) > 1 and "Source Sheet" not in df.columns:
                    df.insert(0, "Source Sheet", sheet_name)
                dfs.append(df)
        except Exception:
            continue

    return merge_dfs(dfs)


# ------------------------------------------------------------
# Text extraction for AI path
# ------------------------------------------------------------
def extract_text_from_pdf(uploaded_file) -> str:
    uploaded_file.seek(0)
    chunks = []

    with pdfplumber.open(uploaded_file) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_parts = []

            try:
                tables = page.extract_tables(
                    table_settings={
                        "vertical_strategy": "lines",
                        "horizontal_strategy": "lines",
                        "intersection_tolerance": 5,
                        "snap_tolerance": 3,
                        "join_tolerance": 3,
                        "edge_min_length": 3,
                        "min_words_vertical": 1,
                        "min_words_horizontal": 1,
                    }
                )
            except Exception:
                tables = []

            useful_tables_found = 0

            for table_index, table in enumerate(tables, start=1):
                if not table:
                    continue

                cleaned_rows = []
                max_cols = 0

                for row in table:
                    if not row:
                        continue

                    cleaned_row = []
                    for cell in row:
                        if cell is None:
                            cleaned_row.append("")
                        else:
                            cleaned_row.append(str(cell).replace("\n", " ").strip())

                    if any(cell != "" for cell in cleaned_row):
                        cleaned_rows.append(cleaned_row)
                        max_cols = max(max_cols, len(cleaned_row))

                if len(cleaned_rows) < 2 or max_cols < 2:
                    continue

                useful_tables_found += 1

                page_parts.append(f"--- PAGE {page_num} | TABLE {table_index} ---")
                for row in cleaned_rows:
                    padded = row + [""] * (max_cols - len(row))
                    page_parts.append(" | ".join(padded))

            page_text = ""
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""

            page_text = page_text.strip()

            if useful_tables_found > 0:
                if page_text:
                    page_parts.append(f"--- PAGE {page_num} | TEXT ---")
                    page_parts.append(page_text)
            else:
                if page_text:
                    page_parts.append(f"--- PAGE {page_num} ---")
                    page_parts.append(page_text)

            if page_parts:
                chunks.append("\n".join(page_parts))

    return "\n\n".join(chunks).strip()


def extract_text_from_docx(uploaded_file) -> str:
    uploaded_file.seek(0)
    doc = Document(uploaded_file)

    chunks = []

    para_texts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if para_texts:
        chunks.append("DOCUMENT TEXT:\n" + "\n".join(para_texts))

    for table_index, table in enumerate(doc.tables, start=1):
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cell for cell in row_cells):
                table_rows.append(" | ".join(row_cells))
        if table_rows:
            chunks.append(f"TABLE {table_index}:\n" + "\n".join(table_rows))

    return "\n\n".join(chunks).strip()


def extract_text_from_xlsx(uploaded_file, selected_sheets=None) -> str:
    uploaded_file.seek(0)
    xls = pd.ExcelFile(uploaded_file)
    chunks = []

    sheet_names = selected_sheets if selected_sheets else xls.sheet_names

    for sheet_name in sheet_names:
        try:
            df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str)
            df = df.fillna("")
            if not df.empty:
                chunks.append(f"SHEET: {sheet_name}")
                chunks.append(df.to_csv(index=False))
        except Exception as exc:
            chunks.append(f"SHEET: {sheet_name}\n[Could not read sheet: {exc}]")

    return "\n\n".join(chunks).strip()


def extract_text_from_csv(uploaded_file) -> str:
    uploaded_file.seek(0)
    raw = uploaded_file.read()

    encodings = ["utf-8", "utf-8-sig", "cp1252", "latin-1"]
    for enc in encodings:
        try:
            return raw.decode(enc)
        except Exception:
            continue

    return raw.decode("utf-8", errors="ignore")


def extract_text_from_json(uploaded_file) -> str:
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    try:
        parsed = json.loads(raw.decode("utf-8"))
        return json.dumps(parsed, ensure_ascii=False, indent=2)
    except Exception:
        return raw.decode("utf-8", errors="ignore")


def extract_text_from_text_like(uploaded_file) -> str:
    uploaded_file.seek(0)
    raw = uploaded_file.read()

    for enc in ["utf-8", "utf-8-sig", "cp1252", "latin-1"]:
        try:
            return raw.decode(enc)
        except Exception:
            pass

    return raw.decode("utf-8", errors="ignore")


def extract_text_for_llm(uploaded_file, selected_sheets=None):
    suffix = Path(uploaded_file.name).suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(uploaded_file), None
    if suffix == ".docx":
        return extract_text_from_docx(uploaded_file), None
    if suffix in [".xlsx", ".xls"]:
        return extract_text_from_xlsx(uploaded_file, selected_sheets=selected_sheets), None
    if suffix == ".csv":
        return extract_text_from_csv(uploaded_file), None
    if suffix == ".json":
        return extract_text_from_json(uploaded_file), None
    if suffix in [".txt", ".md", ".markdown", ".html", ".htm"]:
        return extract_text_from_text_like(uploaded_file), None

    return "", f"Unsupported file type: {suffix}"


# ------------------------------------------------------------
# AI path
# ------------------------------------------------------------
def build_prompt(file_name: str, chunk_text_value: str, chunk_index: int, total_chunks: int) -> str:
    return f"""
Analyze the following extracted content from file "{file_name}".
This is chunk {chunk_index} of {total_chunks}.

Important:
- Some parts may come from PDF table extraction.
- Rows written with "|" separators usually represent table rows and columns.
- Treat those rows as structured table data whenever possible.

Your task:
1. Extract all products/materials/items listed in tables or structured procurement sections.
2. Return ONLY a valid JSON array.
3. Do not include markdown, explanations, headings, or comments.
4. Preserve one object per product/item row.
5. Keep headers / field names in their ORIGINAL language whenever possible.
6. Keep all values in their ORIGINAL language.
7. Do NOT translate anything into English.
8. Preserve the wording from the source as faithfully as possible.
9. If a field is missing, use an empty string.
10. If this chunk has no identifiable product/material table, return [].

CONTENT:
{chunk_text_value}
""".strip()


def call_gemini_once(api_key: str, prompt: str, timeout_seconds: int, min_api_interval_seconds: float):
    enforce_min_api_interval(min_api_interval_seconds)

    url = (
        "https://generativelanguage.googleapis.com/v1/models/"
        f"gemini-2.5-flash:generateContent?key={api_key}"
    )

    payload = {"contents": [{"parts": [{"text": prompt}]}]}
    headers = {"Content-Type": "application/json"}

    response = requests.post(
        url,
        headers=headers,
        json=payload,
        timeout=timeout_seconds,
    )

    st.session_state["last_gemini_call_ts"] = time.time()

    if response.status_code == 429:
        raise requests.exceptions.HTTPError("429 Too Many Requests", response=response)

    response.raise_for_status()
    res_json = response.json()

    try:
        text_response = res_json["candidates"][0]["content"]["parts"][0]["text"]
    except Exception:
        raise RuntimeError(
            f"Unexpected Gemini response: {json.dumps(res_json, ensure_ascii=False)[:2000]}"
        )

    cleaned = clean_llm_json(text_response)

    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"Gemini returned invalid JSON: {exc}\n\nRaw:\n{cleaned[:3000]}")

    if isinstance(data, dict):
        data = [data]

    if not isinstance(data, list):
        raise RuntimeError("Gemini response is not a JSON array.")

    df = pd.DataFrame(data)
    return preserve_original_columns(df)


def call_gemini_with_retry(
    api_key: str,
    prompt: str,
    timeout_seconds: int,
    max_retries: int,
    min_api_interval_seconds: float,
):
    last_error = None

    for attempt in range(max_retries + 1):
        try:
            return call_gemini_once(api_key, prompt, timeout_seconds, min_api_interval_seconds)

        except requests.exceptions.HTTPError as exc:
            status_code = exc.response.status_code if exc.response is not None else None

            if status_code == 429:
                last_error = exc

                retry_after = None
                if exc.response is not None:
                    retry_after = exc.response.headers.get("Retry-After")

                if retry_after:
                    try:
                        sleep_seconds = max(15, int(retry_after))
                    except Exception:
                        sleep_seconds = 30
                else:
                    if attempt == 0:
                        sleep_seconds = random.uniform(20, 30)
                    else:
                        sleep_seconds = random.uniform(45, 60)

                if attempt < max_retries:
                    time.sleep(sleep_seconds)
                    continue

                raise RateLimitError(
                    f"Gemini API rate limit reached (429). Cooldown needed. "
                    f"Last wait was about {sleep_seconds:.1f} seconds."
                )

            if status_code in [500, 502, 503, 504]:
                last_error = exc
                if attempt < max_retries:
                    sleep_seconds = min(30, (2 ** attempt) + random.uniform(0.5, 2.0))
                    time.sleep(sleep_seconds)
                    continue
                raise RuntimeError(f"Gemini server error after retries: HTTP {status_code}")

            raise

        except requests.exceptions.ReadTimeout:
            if attempt < max_retries:
                time.sleep(min(30, (2 ** attempt) + random.uniform(0.5, 2.0)))
                continue
            raise RuntimeError("Gemini request timed out. Try a larger chunk size or retry later.")

        except requests.exceptions.ConnectionError:
            if attempt < max_retries:
                time.sleep(min(30, (2 ** attempt) + random.uniform(0.5, 2.0)))
                continue
            raise RuntimeError("Network connection error while contacting Gemini.")

    raise last_error if last_error else RuntimeError("Gemini call failed.")


def process_file_with_ai(
    uploaded_file,
    api_key: str,
    chunk_size: int,
    timeout_seconds: int,
    max_retries: int,
    min_api_interval_seconds: float,
    selected_sheets=None,
):
    extracted_text, error = extract_text_for_llm(uploaded_file, selected_sheets=selected_sheets)
    if error:
        return None, error, []

    if not extracted_text or not extracted_text.strip():
        return None, "No readable text/content detected in file.", []

    text_chunks = chunk_text(extracted_text, max_chars=chunk_size)
    if not text_chunks:
        return None, "No readable chunks were produced.", []

    per_chunk_logs = []
    dfs = []

    chunk_progress = st.progress(0)
    chunk_status = st.empty()

    for idx, text_chunk in enumerate(text_chunks, start=1):
        chunk_status.info(f"Processing chunk {idx}/{len(text_chunks)} for {uploaded_file.name}")

        prompt = build_prompt(uploaded_file.name, text_chunk, idx, len(text_chunks))

        try:
            df = call_gemini_with_retry(
                api_key=api_key,
                prompt=prompt,
                timeout_seconds=timeout_seconds,
                max_retries=max_retries,
                min_api_interval_seconds=min_api_interval_seconds,
            )
            df = preserve_original_columns(df)
            if not df.empty:
                dfs.append(df)

            per_chunk_logs.append(
                {
                    "chunk": idx,
                    "status": "success",
                    "rows": 0 if df.empty else len(df),
                    "error": "",
                }
            )

        except RateLimitError as exc:
            per_chunk_logs.append(
                {
                    "chunk": idx,
                    "status": "rate_limited",
                    "rows": 0,
                    "error": str(exc),
                }
            )
            break

        except Exception as exc:
            per_chunk_logs.append(
                {
                    "chunk": idx,
                    "status": "error",
                    "rows": 0,
                    "error": str(exc),
                }
            )

        chunk_progress.progress(idx / len(text_chunks))

    if per_chunk_logs and any(log["status"] == "rate_limited" for log in per_chunk_logs):
        return None, "Processing stopped because the Gemini API rate limit was reached.", per_chunk_logs

    final_df = merge_dfs(dfs)
    return final_df, None, per_chunk_logs


# ------------------------------------------------------------
# Main file processor
# ------------------------------------------------------------
def get_selected_sheets_for_file(uploaded_file):
    selection_key = f"sheet_select_{safe_filename(uploaded_file.name)}"
    selected = st.session_state.preview_selections.get(selection_key, [])
    return selected if selected else None


def try_direct_parse(uploaded_file):
    suffix = Path(uploaded_file.name).suffix.lower()

    if suffix == ".csv":
        return direct_parse_csv(uploaded_file), "direct_csv"
    if suffix == ".json":
        return direct_parse_json(uploaded_file), "direct_json"
    if suffix in [".xlsx", ".xls"]:
        selected_sheets = get_selected_sheets_for_file(uploaded_file)
        return direct_parse_excel(uploaded_file, selected_sheets=selected_sheets), "direct_excel"

    return pd.DataFrame(), "not_applicable"


def process_uploaded_file(
    uploaded_file,
    api_key: str,
    chunk_size: int,
    timeout_seconds: int,
    max_retries: int,
    direct_parse_mode: str,
    min_api_interval_seconds: float,
    add_numeric_helpers: bool,
    numeric_threshold: float,
):
    used_method = ""
    chunk_logs = []
    selected_sheets = get_selected_sheets_for_file(uploaded_file)

    if direct_parse_mode == "Use direct parsing first":
        direct_df, used_method = try_direct_parse(uploaded_file)
        if direct_df is not None and not direct_df.empty:
            direct_df = finalize_output_df(
                direct_df,
                add_numeric_helpers=add_numeric_helpers,
                threshold=numeric_threshold,
            )
            return direct_df, None, chunk_logs, used_method

    if not api_key:
        return None, "Gemini API Key is required for AI extraction.", chunk_logs, used_method or "ai_required"

    ai_df, error, chunk_logs = process_file_with_ai(
        uploaded_file=uploaded_file,
        api_key=api_key,
        chunk_size=chunk_size,
        timeout_seconds=timeout_seconds,
        max_retries=max_retries,
        min_api_interval_seconds=min_api_interval_seconds,
        selected_sheets=selected_sheets,
    )

    if ai_df is not None and not ai_df.empty:
        ai_df = finalize_output_df(
            ai_df,
            add_numeric_helpers=add_numeric_helpers,
            threshold=numeric_threshold,
        )

    return ai_df, error, chunk_logs, "ai_extraction"


# ------------------------------------------------------------
# Preview section
# ------------------------------------------------------------
if uploaded_files and show_preview_before_processing:
    render_previews(uploaded_files, preview_rows)


# ------------------------------------------------------------
# Run
# ------------------------------------------------------------
if run_button:
    if not uploaded_files:
        st.error("Please upload at least one file.")
    elif not output_formats:
        st.error("Please select at least one output format.")
    else:
        st.session_state.results = []

        total_files = len(uploaded_files)
        file_progress = st.progress(0)
        file_status = st.empty()

        for file_index, uploaded_file in enumerate(uploaded_files, start=1):
            file_status.info(f"Processing file {file_index}/{total_files}: {uploaded_file.name}")

            try:
                df, error, chunk_logs, method_used = process_uploaded_file(
                    uploaded_file=uploaded_file,
                    api_key=api_key,
                    chunk_size=chunk_size,
                    timeout_seconds=timeout_seconds,
                    max_retries=max_retries,
                    direct_parse_mode=direct_parse_mode,
                    min_api_interval_seconds=min_api_interval_seconds,
                    add_numeric_helpers=add_numeric_helper_columns_flag,
                    numeric_threshold=numeric_detection_threshold,
                )

                if error:
                    st.session_state.results.append(
                        {
                            "file_name": uploaded_file.name,
                            "status": "error",
                            "error": error,
                            "df": None,
                            "chunk_logs": chunk_logs,
                            "method_used": method_used,
                        }
                    )
                else:
                    st.session_state.results.append(
                        {
                            "file_name": uploaded_file.name,
                            "status": "success",
                            "error": None,
                            "df": df if df is not None else pd.DataFrame(),
                            "chunk_logs": chunk_logs,
                            "method_used": method_used,
                        }
                    )

            except Exception as exc:
                st.session_state.results.append(
                    {
                        "file_name": uploaded_file.name,
                        "status": "error",
                        "error": str(exc),
                        "df": None,
                        "chunk_logs": [],
                        "method_used": "unknown",
                    }
                )

            file_progress.progress(file_index / total_files)

            if file_index < total_files and inter_file_delay > 0:
                time.sleep(inter_file_delay)

        file_status.success("All files processed.")


# ------------------------------------------------------------
# Results
# ------------------------------------------------------------
results = st.session_state.get("results", [])

if results:
    st.markdown("## Results")

    success_count = sum(1 for r in results if r["status"] == "success")
    error_count = sum(1 for r in results if r["status"] == "error")

    st.success(f"Completed. Successful: {success_count} | Failed: {error_count}")

    successful_results = [r for r in results if r["status"] == "success"]

    if successful_results:
        try:
            zip_bytes = build_zip(successful_results, output_formats)
            st.download_button(
                label="📦 Download all generated outputs as ZIP",
                data=zip_bytes,
                file_name="all_extracted_outputs.zip",
                mime="application/zip",
            )
        except Exception as exc:
            st.error(f"Could not build ZIP: {exc}")

    mime_map = {
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "csv": "text/csv",
        "html": "text/html",
        "md": "text/markdown",
        "json": "application/json",
    }

    for result in results:
        with st.expander(f"📄 {result['file_name']} — {result['status'].upper()}", expanded=False):
            st.write(f"**Method used:** {result.get('method_used', '')}")

            if result["status"] == "error":
                st.error(result["error"])
            else:
                df = result["df"]

                if df is None or df.empty:
                    st.warning("No extractable items were found.")
                else:
                    st.write(f"Rows extracted: **{len(df)}**")
                    st.dataframe(df, use_container_width=True)

                    cols = st.columns(len(output_formats))
                    for i, fmt in enumerate(output_formats):
                        data = get_export_bytes(df, fmt)
                        cols[i].download_button(
                            label=f"Download {fmt.upper()}",
                            data=data,
                            file_name=f"{safe_filename(result['file_name'])}.{fmt}",
                            mime=mime_map.get(fmt, "application/octet-stream"),
                            key=f"{result['file_name']}_{fmt}",
                        )

            if result.get("chunk_logs"):
                st.markdown("**Chunk log**")
                st.dataframe(pd.DataFrame(result["chunk_logs"]), use_container_width=True)

else:
    st.info("Upload files, preview sheets/tables if needed, choose output formats, and click **Process Files**.")


# ------------------------------------------------------------
# Footer
# ------------------------------------------------------------
st.markdown("---")
st.caption(
    "Powered by Gemini API + direct structured parsing | Multi-file processing | "
    "Language-preserving export | PDF table-first extraction | numeric helper columns | "
    "sheet/table preview before processing | stronger 429 protection | ZIP export"
)
